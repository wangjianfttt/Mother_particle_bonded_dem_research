#!/usr/bin/env python3
"""Build a clean target-neutral submission package from the repaired manuscript."""

from __future__ import annotations

import csv
import hashlib
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "manuscript"
FIG = ROOT / "figures"
SRC = ROOT / "data" / "figure_source"
TABLES = ROOT / "tables"
DOCS = ROOT / "docs"
OUT_ROOT = ROOT / "submission_packages"
SUPPORT_DIR = OUT_ROOT / "repaired_submission_package"
UPLOAD_DIR = OUT_ROOT / "repaired_editorial_upload_ready"
SUPPORT_ZIP = OUT_ROOT / "repaired_submission_package.zip"
UPLOAD_ZIP = OUT_ROOT / "repaired_editorial_upload_ready.zip"

TITLE = (
    "Bonded-template DEM reveals strength- and topology-dependent "
    "fracture-event sequences in packed brittle ceramic pebbles"
)

AUTHORS = [
    ("Jian Wang", "Anhui University of Science and Technology; Institute of Plasma Physics, Chinese Academy of Sciences", "wjfttt@mail.ustc.edu.cn", "Validation, Supervision, Project administration, Methodology, Investigation, Conceptualization"),
    ("Siyu Wang", "Anhui University of Science and Technology", "not provided in current records", "Writing - original draft, Formal analysis, Data curation, Visualization"),
    ("Hang Zhang", "Anhui University of Science and Technology", "not provided in current records", "Writing - original draft, Formal analysis, Data curation, Visualization"),
    ("Ming-Zhun Lei", "Institute of Plasma Physics, Chinese Academy of Sciences", "not provided in current records", "Writing - review and editing, Validation, Project administration, Formal analysis"),
    ("Wei Wen", "Anhui University of Science and Technology; Institute of Plasma Physics, Chinese Academy of Sciences", "not provided in current records", "Writing - review and editing, Supervision, Project administration, Methodology, Conceptualization"),
    ("Qi-Gang Wu", "Institute of Plasma Physics, Chinese Academy of Sciences", "not provided in current records", "Validation, Resources, Data curation"),
    ("Gang Shen", "Anhui University of Science and Technology", "not provided in current records", "Writing - review and editing, Supervision, Project administration"),
    ("Haishun Deng", "Anhui University of Science and Technology", "not provided in current records", "Supervision, Project administration"),
]

FIGURE_STEMS = [
    ("Figure 1", FIG / "main" / "fig1_workflow"),
    ("Figure 2", FIG / "apt_redesign" / "fig2_single_pebble_template_validation"),
    ("Figure 3", FIG / "apt_redesign" / "fig3_entry_state_validation"),
    ("Figure 4", FIG / "apt_redesign" / "fig4_pilot_fracture_event_sequence"),
    ("Figure 5", FIG / "apt_redesign" / "fig5_mechanism_state_space"),
    ("Figure 6", FIG / "pb007" / "pb007_material_strength_response"),
]


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def reset_dir(path: Path) -> None:
    if path.exists():
        shutil.rmtree(path)
    path.mkdir(parents=True, exist_ok=True)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)


def add_title(doc: Document, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string("1F4D78")
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.italic = True
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor.from_string("555555")


def write_markdown_files() -> dict[str, Path]:
    cover = MANUSCRIPT / "repaired_cover_letter_draft.md"
    highlights = MANUSCRIPT / "repaired_highlights.md"
    declaration = MANUSCRIPT / "repaired_declaration_of_competing_interest.md"

    cover.write_text(
        f"""# Cover letter

Dear Editor,

Please consider our manuscript entitled "{TITLE}" as a research article.

This work develops a source-data-backed bonded-template DEM workflow for packed brittle-particle fracture. Each parent particle is represented by a 500-subparticle bonded template, inserted intact into a random load-bearing bed, and compressed while parent-particle fracture events and native force-network descriptors are recorded. The revised manuscript now includes a material-strength matrix across accepted cracking geometries, which directly links bond-strength reduction, local force-path topology, first-event displacement and endpoint damage.

The main contribution is a particle-resolved event-sequence framework for brittle packed beds. It moves beyond bulk force-displacement reporting by identifying the damaged parent particle, the displacement of each localized bond-loss increment, the native force-path state around the event and the response to bond-strength variation. The calculation is intentionally framed as finite-window mechanism evidence for packed brittle ceramic pebbles, with Li4SiO4 used as a representative application material.

All figures in the manuscript are backed by source tables and regeneration scripts. The compact reproducibility package contains the processed event tables, figure source data, scripts, representative DEM inputs, editable figures and a checksum manifest. Very large raw dump and restart files are retained in a separate local/NAS archive and can be supplied on reasonable request subject to repository and transfer limits.

The manuscript has not been published and is not under consideration elsewhere. The authors declare no competing interests.

Sincerely,

Jian Wang  
Corresponding author  
wjfttt@mail.ustc.edu.cn
""",
        encoding="utf-8",
    )

    highlights.write_text(
        """# Highlights

- Intact 500-subparticle bonded templates are inserted into packed beds without pre-loading bond loss.
- Parent-particle fracture events are resolved with displacement, bond-loss increment and damaged-particle identity.
- Native force-network variables link localized microcracking to load-path topology.
- A six-case strength matrix shows geometry-dependent fracture-onset shifts under bond-strength reduction.
- Source tables, editable figures and scripts are packaged for manuscript-level reproducibility.
""",
        encoding="utf-8",
    )

    declaration.write_text(
        """# Declaration of competing interest

The authors declare that they have no known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper.
""",
        encoding="utf-8",
    )
    return {"cover": cover, "highlights": highlights, "declaration": declaration}


def docx_from_markdown(md_path: Path, docx_path: Path, title: str) -> None:
    doc = Document()
    configure_doc(doc)
    add_title(doc, title, TITLE if title != "Cover letter" else None)
    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            if line[2:].strip().lower() != title.lower():
                doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)
    doc.save(docx_path)


def write_author_docx(path: Path) -> None:
    doc = Document()
    configure_doc(doc)
    add_title(doc, "Author e-mails and CRediT contributions", TITLE)
    doc.add_heading("Author Contact Information", level=1)
    p = doc.add_paragraph(
        "Only the corresponding author's e-mail address is currently available in the manuscript metadata. "
        "Fields marked as not provided should be completed before final system submission if the journal requires every author's e-mail address."
    )
    p.runs[0].font.color.rgb = RGBColor.from_string("555555")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Author"
    header[1].text = "Affiliation"
    header[2].text = "E-mail"
    for name, affiliation, email, _credit in AUTHORS:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = affiliation
        cells[2].text = email

    doc.add_heading("CRediT Authorship Contribution Statement", level=1)
    for name, _affiliation, _email, credit in AUTHORS:
        para = doc.add_paragraph()
        para.add_run(f"{name}: ").bold = True
        para.add_run(credit)
    doc.save(path)


def copy_file(src: Path, dst_root: Path, rel: str | None = None, rows: list[dict[str, str]] | None = None, role: str = "support") -> None:
    if not src.exists():
        raise FileNotFoundError(src)
    rel_path = Path(rel) if rel else src.relative_to(ROOT)
    dst = dst_root / rel_path
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst)
    if rows is not None:
        rows.append({"role": role, "path": str(rel_path), "bytes": str(dst.stat().st_size), "sha256": sha256(dst)})


def write_manifest(rows: list[dict[str, str]], path: Path) -> None:
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=["role", "path", "bytes", "sha256"])
        writer.writeheader()
        writer.writerows(rows)


def zip_dir(src_dir: Path, zip_path: Path) -> None:
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(src_dir.rglob("*")):
            if path.is_file():
                zf.write(path, path.relative_to(src_dir.parent))
    zip_path.with_suffix(zip_path.suffix + ".sha256").write_text(f"{sha256(zip_path)}  {zip_path.name}\n", encoding="utf-8")


def build_upload_package(md_files: dict[str, Path]) -> None:
    reset_dir(UPLOAD_DIR)
    docx_from_markdown(md_files["highlights"], UPLOAD_DIR / "02_highlights.docx", "Highlights")
    docx_from_markdown(md_files["declaration"], UPLOAD_DIR / "04_declaration_of_competing_interest.docx", "Declaration of competing interest")
    docx_from_markdown(md_files["cover"], UPLOAD_DIR / "05_cover_letter.docx", "Cover letter")
    write_author_docx(UPLOAD_DIR / "06_author_emails_and_contributions.docx")

    rows: list[dict[str, str]] = []
    upload_files = [
        ("Manuscript", MANUSCRIPT / "repaired_full_submission.pdf", "01_manuscript.pdf"),
        ("Highlights", UPLOAD_DIR / "02_highlights.docx", "02_highlights.docx"),
        ("Graphical abstract", FIG / "apt_redesign" / "apt_graphical_abstract.png", "03_graphical_abstract.png"),
        ("Graphical abstract", FIG / "apt_redesign" / "apt_graphical_abstract.tiff", "03_graphical_abstract.tiff"),
        ("Graphical abstract editable", FIG / "apt_redesign" / "apt_graphical_abstract.pdf", "03_graphical_abstract.pdf"),
        ("Graphical abstract editable", FIG / "apt_redesign" / "apt_graphical_abstract.svg", "03_graphical_abstract.svg"),
        ("Declaration", UPLOAD_DIR / "04_declaration_of_competing_interest.docx", "04_declaration_of_competing_interest.docx"),
        ("Cover letter", UPLOAD_DIR / "05_cover_letter.docx", "05_cover_letter.docx"),
        ("Author details", UPLOAD_DIR / "06_author_emails_and_contributions.docx", "06_author_emails_and_contributions.docx"),
    ]
    # Generated DOCX files already live in UPLOAD_DIR; record them after copying the external files.
    staged_generated = {"02_highlights.docx", "04_declaration_of_competing_interest.docx", "05_cover_letter.docx", "06_author_emails_and_contributions.docx"}
    for role, src, target in upload_files:
        if target not in staged_generated:
            copy_file(src, UPLOAD_DIR, target, rows, role)
        else:
            dst = UPLOAD_DIR / target
            rows.append({"role": role, "path": target, "bytes": str(dst.stat().st_size), "sha256": sha256(dst)})

    readme = UPLOAD_DIR / "README_upload_roles.txt"
    readme.write_text(
        "\n".join(
            [
                "Target-neutral repaired submission upload files",
                "",
                "Manuscript: 01_manuscript.pdf",
                "Highlights: 02_highlights.docx",
                "Graphical abstract: 03_graphical_abstract.png or 03_graphical_abstract.tiff",
                "Graphical abstract editable backup: 03_graphical_abstract.pdf and 03_graphical_abstract.svg",
                "Declaration of competing interest: 04_declaration_of_competing_interest.docx",
                "Cover letter: 05_cover_letter.docx",
                "Author e-mails and CRediT contributions: 06_author_emails_and_contributions.docx",
                "",
                "This upload set is generated from manuscript/repaired_full_submission_draft.md and does not reuse rejected-journal manuscript PDFs.",
            ]
        )
        + "\n",
        encoding="utf-8",
    )
    rows.append({"role": "Upload guide", "path": readme.name, "bytes": str(readme.stat().st_size), "sha256": sha256(readme)})
    write_manifest(rows, UPLOAD_DIR / "MANIFEST.csv")
    zip_dir(UPLOAD_DIR, UPLOAD_ZIP)


def build_support_package(md_files: dict[str, Path]) -> None:
    reset_dir(SUPPORT_DIR)
    rows: list[dict[str, str]] = []

    core_files = [
        MANUSCRIPT / "repaired_full_submission.pdf",
        MANUSCRIPT / "repaired_full_submission.tex",
        MANUSCRIPT / "repaired_full_submission_draft.md",
        MANUSCRIPT / "repaired_full_manuscript_source_data_matrix.csv",
        MANUSCRIPT / "references.bib",
        md_files["cover"],
        md_files["highlights"],
        md_files["declaration"],
        ROOT / "START_HERE_CPM_SUBMISSION.md",
        ROOT / "README_CPM_SUBMISSION_20260704.md",
        DOCS / "repaired_full_pdf_visual_qa_20260704.md",
        DOCS / "nas_raw_dump_storage_check_20260704_1736.md",
        DOCS / "next_stage_optimization_plan.md",
        DOCS / "cpm_live_submission_packet_docx_qa_20260704.md",
        DOCS / "cpm_author_email_public_lookup_20260704.md",
        DOCS / "cpm_author_email_public_lookup_20260704.csv",
        DOCS / "cpm_goal_completion_audit_20260704.md",
        DOCS / "cpm_goal_completion_audit_20260704.csv",
        DOCS / "cpm_goal_completion_audit_20260704.json",
        MANUSCRIPT / "computational_particle_mechanics_live_submission_packet.docx",
        MANUSCRIPT / "computational_particle_mechanics_author_email_collection_packet.docx",
        MANUSCRIPT / "computational_particle_mechanics_author_email_collection_packet.md",
        MANUSCRIPT / "computational_particle_mechanics_author_email_collection_packet.txt",
        MANUSCRIPT / "computational_particle_mechanics_author_email_collection_packet.csv",
        MANUSCRIPT / "computational_particle_mechanics_coauthor_email_request_zh_en.docx",
        MANUSCRIPT / "computational_particle_mechanics_coauthor_email_request_zh_en.txt",
        MANUSCRIPT / "computational_particle_mechanics_live_submission_checklist.docx",
        MANUSCRIPT / "computational_particle_mechanics_live_submission_checklist.md",
    ]
    for src in core_files:
        copy_file(src, SUPPORT_DIR, rows=rows, role="core")

    for name in ["02_highlights.docx", "04_declaration_of_competing_interest.docx", "05_cover_letter.docx", "06_author_emails_and_contributions.docx"]:
        copy_file(UPLOAD_DIR / name, SUPPORT_DIR, f"manuscript/{name}", rows, "submission docx")

    for _label, stem in FIGURE_STEMS:
        for ext in [".pdf", ".svg", ".png", ".tiff"]:
            copy_file(stem.with_suffix(ext), SUPPORT_DIR, rows=rows, role="figure")
    for ext in [".pdf", ".svg", ".png", ".tiff"]:
        copy_file((FIG / "apt_redesign" / "apt_graphical_abstract").with_suffix(ext), SUPPORT_DIR, rows=rows, role="graphical abstract")

    source_files = set()
    with (MANUSCRIPT / "repaired_full_manuscript_source_data_matrix.csv").open(encoding="utf-8") as handle:
        reader = csv.DictReader(handle)
        for row in reader:
            for field in ["source_data", "regeneration_or_check_script"]:
                for item in row[field].split(";"):
                    item = item.strip()
                    if item and not item.startswith("manuscript/repaired_full_submission_draft.md"):
                        source_files.add(item)
    extra_sources = [
        "scripts/build_repaired_full_latex.py",
        "scripts/build_repaired_submission_package.py",
        "scripts/check_repaired_submission_package.py",
        "scripts/check_repaired_full_manuscript_consistency.py",
        "scripts/build_apt_redesigned_data_figures.py",
        "scripts/plot_apt_graphical_abstract.py",
        "scripts/build_cpm_literature_gap_map.py",
        "scripts/build_cpm_official_submission_guide_alignment.py",
        "scripts/build_cpm_blinded_review_package.py",
        "scripts/build_cpm_material_response_summary.py",
        "scripts/build_cpm_reviewer_risk_preflight.py",
        "scripts/build_cpm_submission_readiness_report.py",
        "scripts/build_cpm_live_submission_packet.py",
        "scripts/build_cpm_submission_support_docs.py",
        "scripts/build_cpm_goal_completion_audit.py",
        "scripts/check_cpm_reviewer_risk_preflight.py",
        "scripts/check_cpm_scientific_alignment.py",
        "data/figure_source/pb007_material_strength_response.csv",
        "tables/pb007_material_parameter_response.csv",
        "tables/pb007_material_strength_matrix_summary.csv",
        "tables/pb007_material_parameter_run_progress.csv",
        "docs/cpm_literature_gap_map_20260704.csv",
        "docs/cpm_literature_gap_map_20260704.md",
        "docs/cpm_official_submission_guide_alignment_20260704.csv",
        "docs/cpm_official_submission_guide_alignment_20260704.md",
        "docs/cpm_material_response_summary_20260704.csv",
        "docs/cpm_material_response_summary_20260704.md",
        "docs/cpm_reviewer_risk_preflight_20260704.csv",
        "docs/cpm_reviewer_risk_preflight_20260704.md",
        "docs/cpm_live_submission_packet_20260704.csv",
        "docs/cpm_live_submission_packet_20260704.json",
        "docs/cpm_live_submission_packet_20260704.md",
        "docs/cpm_author_email_public_lookup_20260704.md",
        "docs/cpm_author_email_public_lookup_20260704.csv",
        "docs/cpm_goal_completion_audit_20260704.md",
        "docs/cpm_goal_completion_audit_20260704.csv",
        "docs/cpm_goal_completion_audit_20260704.json",
    ]
    source_files.update(extra_sources)
    for rel in sorted(source_files):
        path = ROOT / rel
        if path.exists() and path.is_file():
            copy_file(path, SUPPORT_DIR, rows=rows, role="source/script")

    readme = SUPPORT_DIR / "README.md"
    readme.write_text(
        f"""# Repaired submission support package

This package supports the manuscript:

**{TITLE}**

The package is generated from the repaired target-neutral manuscript, not from the older APT/JNM upload bundles. It contains the manuscript PDF/TeX/Markdown source, submission-side Word documents, source-data matrices, editable figures, raster figures, figure source tables and regeneration/check scripts.

Repository DOI: https://doi.org/10.5281/zenodo.20687351

Large raw DEM dump, local-bond and restart histories are kept outside this compact package on the NAS archive described in `docs/nas_raw_dump_storage_check_20260704_1736.md`.

The included `docs/cpm_*` files record the Computational Particle Mechanics retargeting support evidence, including official-guide alignment, literature-gap mapping, material-response summary values, reviewer-risk preflight checks, the live-submission packet and public candidate e-mail lookup for coauthor confirmation. The live readiness report is kept outside this zip because it records the current package checksum.

Primary checks:

```bash
python3 scripts/check_repaired_full_manuscript_consistency.py
python3 scripts/build_apt_redesigned_data_figures.py
python3 scripts/build_cpm_official_submission_guide_alignment.py
python3 scripts/build_cpm_material_response_summary.py
python3 scripts/build_cpm_reviewer_risk_preflight.py
python3 scripts/check_cpm_reviewer_risk_preflight.py
python3 scripts/build_repaired_full_latex.py
(cd manuscript && latexmk -pdf -interaction=nonstopmode -halt-on-error repaired_full_submission.tex)
```
""",
        encoding="utf-8",
    )
    rows.append({"role": "Readme", "path": "README.md", "bytes": str(readme.stat().st_size), "sha256": sha256(readme)})
    write_manifest(rows, SUPPORT_DIR / "MANIFEST.csv")
    zip_dir(SUPPORT_DIR, SUPPORT_ZIP)


def main() -> None:
    md_files = write_markdown_files()
    build_upload_package(md_files)
    build_support_package(md_files)
    print(UPLOAD_DIR)
    print(UPLOAD_ZIP)
    print(UPLOAD_ZIP.with_suffix(UPLOAD_ZIP.suffix + ".sha256"))
    print(SUPPORT_DIR)
    print(SUPPORT_ZIP)
    print(SUPPORT_ZIP.with_suffix(SUPPORT_ZIP.suffix + ".sha256"))


if __name__ == "__main__":
    main()

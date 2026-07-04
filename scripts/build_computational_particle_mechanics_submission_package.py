#!/usr/bin/env python3
"""Build a Computational Particle Mechanics upload package from repaired files."""

from __future__ import annotations

import csv
import hashlib
import subprocess
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "manuscript"
FIG = ROOT / "figures"
OUT_ROOT = ROOT / "submission_packages"
UPLOAD_DIR = OUT_ROOT / "computational_particle_mechanics_upload_ready"
UPLOAD_ZIP = OUT_ROOT / "computational_particle_mechanics_upload_ready.zip"
CPM_TEX = MANUSCRIPT / "computational_particle_mechanics_submission.tex"
CPM_PDF = MANUSCRIPT / "computational_particle_mechanics_submission.pdf"

TITLE = (
    "Bonded-template DEM reveals strength- and topology-dependent "
    "fracture-event sequences in packed brittle ceramic pebbles"
)

AUTHORS = [
    ("Jian Wang", "Anhui University of Science and Technology; Institute of Plasma Physics, Chinese Academy of Sciences", "wjfttt@mail.ustc.edu.cn", "Validation, Supervision, Project administration, Methodology, Investigation, Conceptualization"),
    ("Siyu Wang", "Anhui University of Science and Technology", "not provided in current records", "Writing - original draft, Formal analysis, Data curation, Visualization"),
    ("Hang Zhang", "Anhui University of Science and Technology", "not provided in current records", "Writing - original draft, Formal analysis, Data curation, Visualization"),
    ("Ming-Zhun Lei", "Institute of Plasma Physics, Chinese Academy of Sciences", "not provided in current records", "Writing - review and editing, Validation, Project administration, Formal analysis"),
    ("Wei Wen", "Anhui University of Science and Technology; Institute of Plasma Physics, Chinese Academy of Sciences", "not provided in current records", "Writing - review and editing, Supervision, Project administration, Methodology, Conceptualization"),
    ("Qi-Gang Wu", "Institute of Plasma Physics, Chinese Academy of Sciences", "not provided in current records", "Validation, Resources, Data curation"),
    ("Gang Shen", "Anhui University of Science and Technology", "not provided in current records", "Writing - review and editing, Supervision, Project administration"),
    ("Haishun Deng", "Anhui University of Science and Technology", "not provided in current records", "Supervision, Project administration"),
]


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def reset_dir(path: Path) -> None:
    if path.exists():
        shutil.rmtree(path)
    path.mkdir(parents=True, exist_ok=True)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size in [("Heading 1", 16), ("Heading 2", 13)]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string("1F4D78")


def add_title(doc: Document, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string("1F4D78")
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.italic = True
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor.from_string("555555")


def docx_from_markdown(md_path: Path, docx_path: Path, title: str) -> None:
    doc = Document()
    configure_doc(doc)
    add_title(doc, title, TITLE if title != "Cover letter" else None)
    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            if line[2:].strip().lower() != title.lower():
                doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)
    doc.save(docx_path)


def write_author_docx(path: Path) -> None:
    doc = Document()
    configure_doc(doc)
    add_title(doc, "Author e-mails and CRediT contributions", TITLE)
    doc.add_heading("Author Contact Information", level=1)
    doc.add_paragraph(
        "Only the corresponding author's e-mail address is currently available in the manuscript metadata. "
        "Fields marked as not provided should be completed before final system submission if the journal requires every author's e-mail address."
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Author"
    header[1].text = "Affiliation"
    header[2].text = "E-mail"
    for name, affiliation, email, _credit in AUTHORS:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = affiliation
        cells[2].text = email
    doc.add_heading("CRediT Authorship Contribution Statement", level=1)
    for name, _affiliation, _email, credit in AUTHORS:
        para = doc.add_paragraph()
        para.add_run(f"{name}: ").bold = True
        para.add_run(credit)
    doc.save(path)


def copy_and_record(src: Path, dst: Path, role: str, rows: list[dict[str, str]]) -> None:
    if not src.exists():
        raise FileNotFoundError(src)
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst)
    rows.append(
        {
            "role": role,
            "path": str(dst.relative_to(UPLOAD_DIR)),
            "bytes": str(dst.stat().st_size),
            "sha256": sha256(dst),
        }
    )


def zip_dir(src_dir: Path, zip_path: Path) -> None:
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(src_dir.rglob("*")):
            if path.is_file():
                zf.write(path, path.relative_to(src_dir.parent))
    zip_path.with_suffix(zip_path.suffix + ".sha256").write_text(
        f"{sha256(zip_path)}  {zip_path.name}\n", encoding="utf-8"
    )


def build_latex_source_zip(path: Path) -> None:
    if path.exists():
        path.unlink()
    readme = (
        "Computational Particle Mechanics LaTeX source package\n\n"
        "Compile from the manuscript directory after extraction:\n\n"
        "  cd manuscript\n"
        "  latexmk -pdf -interaction=nonstopmode -halt-on-error computational_particle_mechanics_submission.tex\n\n"
        "The figure paths are relative to the manuscript directory and point to ../figures/.\n"
    )
    latex_files = [
        CPM_TEX,
        MANUSCRIPT / "references.bib",
    ]
    figure_files = [
        FIG / "main" / "fig1_workflow.pdf",
        FIG / "apt_redesign" / "fig2_single_pebble_template_validation.pdf",
        FIG / "apt_redesign" / "fig3_entry_state_validation.pdf",
        FIG / "apt_redesign" / "fig4_pilot_fracture_event_sequence.pdf",
        FIG / "apt_redesign" / "fig5_mechanism_state_space.pdf",
        FIG / "pb007" / "pb007_material_strength_response.pdf",
    ]
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("README_latex_source.txt", readme)
        for src in latex_files:
            if not src.exists():
                raise FileNotFoundError(src)
            zf.write(src, src.relative_to(ROOT))
        for src in figure_files:
            if not src.exists():
                raise FileNotFoundError(src)
            zf.write(src, src.relative_to(ROOT))


def build_main_figures_zip(path: Path) -> None:
    """Package main figures as separate editable/raster files for upload."""
    if path.exists():
        path.unlink()
    figure_stems = [
        FIG / "main" / "fig1_workflow",
        FIG / "apt_redesign" / "fig2_single_pebble_template_validation",
        FIG / "apt_redesign" / "fig3_entry_state_validation",
        FIG / "apt_redesign" / "fig4_pilot_fracture_event_sequence",
        FIG / "apt_redesign" / "fig5_mechanism_state_space",
        FIG / "pb007" / "pb007_material_strength_response",
    ]
    readme = (
        "Main figure files for Computational Particle Mechanics submission\n\n"
        "Each main figure is supplied as PDF, PNG and SVG. The high-resolution TIFF\n"
        "files remain in the local working tree and can be uploaded separately if\n"
        "the submission system requests TIFF-only figure files.\n"
    )
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("README_main_figures.txt", readme)
        for stem in figure_stems:
            for ext in [".pdf", ".png", ".svg"]:
                src = stem.with_suffix(ext)
                if not src.exists():
                    raise FileNotFoundError(src)
                zf.write(src, src.relative_to(ROOT))


def build_target_tex_and_pdf() -> None:
    source = MANUSCRIPT / "repaired_full_submission.tex"
    text = source.read_text(encoding="utf-8")
    text = text.replace(r"\journal{To be determined}", r"\journal{Computational Particle Mechanics}")
    CPM_TEX.write_text(text, encoding="utf-8")
    subprocess.run(
        [
            "latexmk",
            "-pdf",
            "-interaction=nonstopmode",
            "-halt-on-error",
            CPM_TEX.name,
        ],
        cwd=MANUSCRIPT,
        check=True,
    )


def main() -> None:
    build_target_tex_and_pdf()
    reset_dir(UPLOAD_DIR)
    rows: list[dict[str, str]] = []

    cover_docx = UPLOAD_DIR / "05_cover_letter.docx"
    highlights_docx = UPLOAD_DIR / "02_highlights.docx"
    declaration_docx = UPLOAD_DIR / "04_declaration_of_competing_interest.docx"
    author_docx = UPLOAD_DIR / "06_author_emails_and_contributions.docx"
    fields_docx = UPLOAD_DIR / "08_editorial_submission_fields.docx"

    docx_from_markdown(MANUSCRIPT / "computational_particle_mechanics_cover_letter.md", cover_docx, "Cover letter")
    docx_from_markdown(MANUSCRIPT / "computational_particle_mechanics_highlights.md", highlights_docx, "Highlights")
    docx_from_markdown(MANUSCRIPT / "repaired_declaration_of_competing_interest.md", declaration_docx, "Declaration of competing interest")
    docx_from_markdown(MANUSCRIPT / "computational_particle_mechanics_editorial_fields.md", fields_docx, "Computational Particle Mechanics Editorial Submission Fields")
    write_author_docx(author_docx)
    latex_source_zip = UPLOAD_DIR / "07_latex_source.zip"
    build_latex_source_zip(latex_source_zip)
    main_figures_zip = UPLOAD_DIR / "09_main_figures.zip"
    build_main_figures_zip(main_figures_zip)

    copy_and_record(CPM_PDF, UPLOAD_DIR / "01_manuscript.pdf", "Manuscript", rows)
    for path, role in [
        (highlights_docx, "Highlights"),
        (declaration_docx, "Declaration"),
        (cover_docx, "Cover letter"),
        (author_docx, "Author details"),
        (latex_source_zip, "LaTeX source"),
        (fields_docx, "Editorial submission fields"),
        (main_figures_zip, "Main figure files"),
    ]:
        rows.append({"role": role, "path": path.name, "bytes": str(path.stat().st_size), "sha256": sha256(path)})

    for ext, role in [
        (".png", "Graphical abstract"),
        (".tiff", "Graphical abstract"),
        (".pdf", "Graphical abstract editable"),
        (".svg", "Graphical abstract editable"),
    ]:
        copy_and_record(
            (FIG / "apt_redesign" / "apt_graphical_abstract").with_suffix(ext),
            UPLOAD_DIR / f"03_graphical_abstract{ext}",
            role,
            rows,
        )

    readme = UPLOAD_DIR / "README_upload_roles.txt"
    readme.write_text(
        "\n".join(
            [
                "Computational Particle Mechanics upload files",
                "",
                "Manuscript: 01_manuscript.pdf",
                "Highlights: 02_highlights.docx",
                "Graphical abstract: 03_graphical_abstract.png or 03_graphical_abstract.tiff",
                "Editable graphical abstract backup: 03_graphical_abstract.pdf and 03_graphical_abstract.svg",
                "Declaration of competing interest: 04_declaration_of_competing_interest.docx",
                "Cover letter: 05_cover_letter.docx",
                "Author e-mails and CRediT contributions: 06_author_emails_and_contributions.docx",
                "LaTeX manuscript source: 07_latex_source.zip",
                "Editorial system paste fields: 08_editorial_submission_fields.docx",
                "Main figure files: 09_main_figures.zip",
                "",
                "Computational Particle Mechanics currently routes submissions through the Elsevier/ScienceDirect journal page.",
                "The full reproducibility package remains submission_packages/repaired_submission_package.zip.",
            ]
        )
        + "\n",
        encoding="utf-8",
    )
    rows.append({"role": "Upload guide", "path": readme.name, "bytes": str(readme.stat().st_size), "sha256": sha256(readme)})

    manifest = UPLOAD_DIR / "MANIFEST.csv"
    with manifest.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=["role", "path", "bytes", "sha256"])
        writer.writeheader()
        writer.writerows(rows)
    zip_dir(UPLOAD_DIR, UPLOAD_ZIP)
    print(UPLOAD_DIR)
    print(UPLOAD_ZIP)
    print(UPLOAD_ZIP.with_suffix(UPLOAD_ZIP.suffix + ".sha256"))


if __name__ == "__main__":
    main()

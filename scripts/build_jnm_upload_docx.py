#!/usr/bin/env python3
"""Build DOCX upload companions for the JNM cover letter and declarations."""

from __future__ import annotations

import re
import shutil
import subprocess
import sys
import os
from pathlib import Path

BUNDLED_PYTHON = Path.home() / ".cache/codex-runtimes/codex-primary-runtime/dependencies/python/bin/python3"
if BUNDLED_PYTHON.exists() and Path(sys.executable).resolve() != BUNDLED_PYTHON.resolve():
    os.execv(str(BUNDLED_PYTHON), [str(BUNDLED_PYTHON), *sys.argv])

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
COVER_MD = ROOT / "manuscript/journal_of_nuclear_materials_cover_letter_draft.md"
DECL_MD = ROOT / "manuscript/journal_of_nuclear_materials_elsevier_declarations.md"
COVER_DOCX = ROOT / "manuscript/journal_of_nuclear_materials_cover_letter.docx"
DECL_DOCX = ROOT / "manuscript/journal_of_nuclear_materials_elsevier_declarations.docx"
QA_REPORT = ROOT / "docs/jnm_upload_docx_qa.md"
RENDER_SCRIPT = (
    Path.home()
    / ".codex/plugins/cache/openai-primary-runtime/documents/26.601.10930/skills/documents/render_docx.py"
)


def setup_document(title: str, subtitle: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)

    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(31, 77, 120)
    title_para.paragraph_format.space_after = Pt(3)

    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run(subtitle)
    subtitle_run.italic = True
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(85, 85, 85)
    subtitle_para.paragraph_format.space_after = Pt(12)
    return doc


def add_markdown_text(doc: Document, text: str) -> None:
    for raw in text.splitlines():
        line = raw.rstrip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("- [ ] "):
            para = doc.add_paragraph(style=None)
            para.style = doc.styles["List Bullet"]
            para.add_run("[ ] " + line[6:].strip())
            continue
        if line.startswith("- "):
            para = doc.add_paragraph(style=None)
            para.style = doc.styles["List Bullet"]
            para.add_run(line[2:].strip())
            continue

        para = doc.add_paragraph()
        parts = re.split(r"(`[^`]+`)", line)
        for part in parts:
            if part.startswith("`") and part.endswith("`"):
                run = para.add_run(part[1:-1])
                run.font.name = "Courier New"
                run.font.size = Pt(10)
            else:
                # Preserve the Markdown hard line-break convention used in the signature block.
                for idx, subpart in enumerate(part.split("  ")):
                    if idx:
                        para.add_run().add_break(WD_BREAK.LINE)
                    if subpart:
                        para.add_run(subpart)


def write_cover_letter() -> None:
    doc = setup_document(
        "Cover letter",
        "Journal of Nuclear Materials submission companion document",
    )
    add_markdown_text(doc, COVER_MD.read_text(encoding="utf-8"))
    doc.save(COVER_DOCX)


def write_declarations() -> None:
    doc = setup_document(
        "Declarations",
        "Journal of Nuclear Materials Editorial Manager companion document",
    )
    add_markdown_text(doc, DECL_MD.read_text(encoding="utf-8"))
    doc.save(DECL_DOCX)


def render_check(path: Path) -> str:
    if shutil.which("soffice") is None and shutil.which("libreoffice") is None:
        return "SKIPPED: LibreOffice/soffice is not available in this environment."
    outdir = ROOT / "docs/upload_docx_render_qa" / path.stem
    outdir.mkdir(parents=True, exist_ok=True)
    result = subprocess.run(
        ["python3", str(RENDER_SCRIPT), str(path), "--output_dir", str(outdir), "--emit_pdf"],
        cwd=ROOT,
        text=True,
        capture_output=True,
        check=False,
    )
    if result.returncode != 0:
        return "FAILED: " + (result.stdout + result.stderr).strip()
    return f"PASS: rendered to {outdir.relative_to(ROOT)}"


def main() -> int:
    write_cover_letter()
    write_declarations()
    cover_render = render_check(COVER_DOCX)
    decl_render = render_check(DECL_DOCX)
    lines = [
        "# JNM upload DOCX QA",
        "",
        "Generated DOCX companion files for Editorial Manager upload/paste support.",
        "",
        f"- Cover letter DOCX: `{COVER_DOCX.relative_to(ROOT)}`",
        f"- Declarations DOCX: `{DECL_DOCX.relative_to(ROOT)}`",
        "",
        "## Render QA",
        "",
        f"- Cover letter: {cover_render}",
        f"- Declarations: {decl_render}",
        "",
        "## Structural QA",
        "",
        "- DOCX files are generated from the active Markdown source files.",
        "- The selected document preset is `standard_business_brief`: Letter page, 1 inch margins, Calibri 11 pt body, restrained blue heading hierarchy.",
        "- If visual render QA is skipped, the corresponding author should visually inspect the DOCX files in Word, LibreOffice or WPS before final upload.",
        "",
    ]
    QA_REPORT.write_text("\n".join(lines), encoding="utf-8")
    print(COVER_DOCX)
    print(DECL_DOCX)
    print(QA_REPORT)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

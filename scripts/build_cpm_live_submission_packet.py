#!/usr/bin/env python3
"""Build a live-submission packet for CPM online submission."""

from __future__ import annotations

import csv
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
MANUSCRIPT = ROOT / "manuscript"
UPLOAD_DIR = ROOT / "submission_packages" / "computational_particle_mechanics_upload_ready"

FIELDS = MANUSCRIPT / "computational_particle_mechanics_editorial_fields.md"
EMAIL_SHEET = UPLOAD_DIR / "10_author_email_completion_sheet.csv"
MANIFEST = UPLOAD_DIR / "MANIFEST.csv"

OUT_MD = DOCS / "cpm_live_submission_packet_20260704.md"
OUT_CSV = DOCS / "cpm_live_submission_packet_20260704.csv"
OUT_JSON = DOCS / "cpm_live_submission_packet_20260704.json"
OUT_DOCX = MANUSCRIPT / "computational_particle_mechanics_live_submission_packet.docx"

TITLE = "Bonded-template DEM reveals strength- and topology-dependent fracture-event sequences in packed brittle ceramic pebbles"
TARGET = "Computational Particle Mechanics"
ROUTE = "ScienceDirect / Editorial Manager live submission route"


def read_sections(path: Path) -> dict[str, str]:
    sections: dict[str, list[str]] = {}
    current: str | None = None
    for line in path.read_text(encoding="utf-8").splitlines():
        if line.startswith("## "):
            current = line[3:].strip()
            sections[current] = []
        elif current:
            sections[current].append(line.rstrip())
    return {key: "\n".join(value).strip() for key, value in sections.items()}


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def upload_steps() -> list[dict[str, str]]:
    return [
        {"step": "1", "item": "Manuscript", "file_or_value": "01_manuscript.pdf", "required": "yes", "note": "Use the full manuscript unless the live system asks for a blinded file."},
        {"step": "2", "item": "Highlights", "file_or_value": "02_highlights.docx", "required": "yes", "note": "Five highlights are also available in the editorial fields document."},
        {"step": "3", "item": "Graphical abstract", "file_or_value": "03_graphical_abstract.png or 03_graphical_abstract.tiff", "required": "if requested", "note": "PDF and SVG backups are included for editable artwork."},
        {"step": "4", "item": "Declaration of competing interest", "file_or_value": "04_declaration_of_competing_interest.docx", "required": "yes", "note": "No competing interests declared."},
        {"step": "5", "item": "Cover letter", "file_or_value": "05_cover_letter.docx", "required": "yes", "note": "Targeted to Computational Particle Mechanics."},
        {"step": "6", "item": "Author details and contributions", "file_or_value": "06_author_emails_and_contributions.docx", "required": "yes", "note": "Use with completed coauthor e-mail sheet."},
        {"step": "7", "item": "LaTeX source", "file_or_value": "07_latex_source.zip", "required": "if requested", "note": "Editable source for the full manuscript."},
        {"step": "8", "item": "Editorial paste fields", "file_or_value": "08_editorial_submission_fields.docx", "required": "support", "note": "Use for title, abstract, keywords, data and code availability."},
        {"step": "9", "item": "Main figures", "file_or_value": "09_main_figures.zip", "required": "if requested", "note": "Contains separate PDF, PNG and SVG figure files."},
        {"step": "10", "item": "Author e-mail completion sheet", "file_or_value": "10_author_email_completion_sheet.docx / .csv", "required": "external metadata", "note": "Seven coauthor e-mails remain missing in local records."},
        {"step": "11", "item": "Optional blinded manuscript", "file_or_value": "computational_particle_mechanics_blinded_review_optional.zip", "required": "only if requested", "note": "Use for double-anonymized review workflow if the system asks for a blinded file."},
        {"step": "12", "item": "Reduced reproducibility package", "file_or_value": "repaired_submission_package.zip", "required": "support", "note": "Use for data/code repository or supplemental support if requested."},
    ]


def paste_fields(sections: dict[str, str]) -> list[dict[str, str]]:
    return [
        {"field": "Journal", "value": TARGET, "source": "START_HERE_CPM_SUBMISSION.md"},
        {"field": "Submission route", "value": ROUTE, "source": "docs/cpm_official_submission_guide_alignment_20260704.md"},
        {"field": "Article type", "value": sections["Article type"], "source": str(FIELDS.relative_to(ROOT))},
        {"field": "Title", "value": sections["Manuscript title"], "source": str(FIELDS.relative_to(ROOT))},
        {"field": "Keywords", "value": sections["Keywords"], "source": str(FIELDS.relative_to(ROOT))},
        {"field": "Abstract", "value": sections["Abstract"], "source": str(FIELDS.relative_to(ROOT))},
        {"field": "Highlights", "value": sections["Highlights"], "source": str(FIELDS.relative_to(ROOT))},
        {"field": "Data availability", "value": sections["Data availability statement"], "source": str(FIELDS.relative_to(ROOT))},
        {"field": "Code availability", "value": sections["Code availability statement"], "source": str(FIELDS.relative_to(ROOT))},
        {"field": "Declaration", "value": sections["Declaration of competing interest"], "source": str(FIELDS.relative_to(ROOT))},
        {"field": "Funding / acknowledgements", "value": sections["Funding / acknowledgements text"], "source": str(FIELDS.relative_to(ROOT))},
        {"field": "Corresponding author", "value": sections["Corresponding author"], "source": str(FIELDS.relative_to(ROOT))},
    ]


def build_payload() -> dict[str, object]:
    sections = read_sections(FIELDS)
    manifest_rows = read_csv(MANIFEST)
    author_rows = read_csv(EMAIL_SHEET)
    missing = [row["Author"] for row in author_rows if row["Status"] == "Missing"]
    return {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "target_journal": TARGET,
        "submission_route": ROUTE,
        "title": TITLE,
        "upload_package": "submission_packages/computational_particle_mechanics_upload_ready.zip",
        "optional_blinded_package": "submission_packages/computational_particle_mechanics_blinded_review_optional.zip",
        "reduced_reproducibility_package": "submission_packages/repaired_submission_package.zip",
        "upload_steps": upload_steps(),
        "paste_fields": paste_fields(sections),
        "manifest_rows": manifest_rows,
        "author_rows": author_rows,
        "missing_email_count": len(missing),
        "missing_email_authors": missing,
        "final_external_actions": [
            "Fill seven coauthor e-mail addresses if required by the live submission system.",
            "Confirm article type/category in the live system.",
            "Preview the system-generated PDF before final submit.",
        ],
    }


def write_csv_packet(payload: dict[str, object]) -> None:
    rows = payload["upload_steps"]
    assert isinstance(rows, list)
    with OUT_CSV.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=["step", "item", "file_or_value", "required", "note"], lineterminator="\n")
        writer.writeheader()
        writer.writerows(rows)  # type: ignore[arg-type]


def write_markdown(payload: dict[str, object]) -> None:
    lines = [
        "# CPM Live Submission Packet",
        "",
        f"Generated: `{payload['generated_at']}`",
        "",
        f"- Target journal: {payload['target_journal']}",
        f"- Submission route: {payload['submission_route']}",
        f"- Manuscript title: {payload['title']}",
        f"- Main upload package: `{payload['upload_package']}`",
        f"- Optional blinded package: `{payload['optional_blinded_package']}`",
        f"- Reduced reproducibility package: `{payload['reduced_reproducibility_package']}`",
        "",
        "## Upload Sequence",
        "",
        "| Step | Item | File or value | Required | Note |",
        "| ---: | --- | --- | --- | --- |",
    ]
    for row in payload["upload_steps"]:  # type: ignore[index]
        lines.append(
            f"| {row['step']} | {row['item']} | `{row['file_or_value']}` | {row['required']} | {row['note']} |"
        )
    lines.extend(["", "## Paste Fields", ""])
    for row in payload["paste_fields"]:  # type: ignore[index]
        lines.append(f"### {row['field']}")
        lines.append("")
        lines.append(str(row["value"]))
        lines.append("")
        lines.append(f"Source: `{row['source']}`")
        lines.append("")
    lines.extend(["## Missing Author E-mails", ""])
    lines.append(f"Missing count: `{payload['missing_email_count']}`")
    lines.append("")
    for name in payload["missing_email_authors"]:  # type: ignore[index]
        lines.append(f"- {name}")
    lines.extend(["", "## Final External Actions", ""])
    for item in payload["final_external_actions"]:  # type: ignore[index]
        lines.append(f"- [ ] {item}")
    lines.append("")
    OUT_MD.write_text("\n".join(lines), encoding="utf-8")


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].paragraph_format.space_after = Pt(4)
    for name, size in [("Heading 1", 14), ("Heading 2", 11.5), ("Heading 3", 10.5)]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 78, 121)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(8.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)


def write_docx(payload: dict[str, object]) -> None:
    doc = Document()
    configure_doc(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CPM Live Submission Packet")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 78, 121)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(str(payload["title"]))

    doc.add_heading("Core Package", level=1)
    core_rows = [
        ["Target journal", str(payload["target_journal"])],
        ["Submission route", str(payload["submission_route"])],
        ["Main upload package", str(payload["upload_package"])],
        ["Optional blinded package", str(payload["optional_blinded_package"])],
        ["Reduced reproducibility package", str(payload["reduced_reproducibility_package"])],
        ["Missing author e-mails", str(payload["missing_email_count"])],
    ]
    add_table(doc, ["Field", "Value"], core_rows)

    doc.add_heading("Upload Sequence", level=1)
    upload_rows = [
        [row["step"], row["item"], row["file_or_value"], row["required"], row["note"]]
        for row in payload["upload_steps"]  # type: ignore[index]
    ]
    add_table(doc, ["Step", "Item", "File or value", "Required", "Note"], upload_rows)

    doc.add_heading("Paste Fields", level=1)
    for row in payload["paste_fields"]:  # type: ignore[index]
        doc.add_heading(row["field"], level=2)
        doc.add_paragraph(str(row["value"]))
        note = doc.add_paragraph()
        note.add_run("Source: ").bold = True
        note.add_run(str(row["source"]))

    doc.add_heading("Missing Author E-mails", level=1)
    for name in payload["missing_email_authors"]:  # type: ignore[index]
        doc.add_paragraph(str(name), style="List Bullet")

    doc.add_heading("Final External Actions", level=1)
    for item in payload["final_external_actions"]:  # type: ignore[index]
        doc.add_paragraph(str(item), style="List Bullet")
    doc.save(OUT_DOCX)


def main() -> None:
    DOCS.mkdir(exist_ok=True)
    MANUSCRIPT.mkdir(exist_ok=True)
    payload = build_payload()
    OUT_JSON.write_text(json.dumps(payload, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    write_csv_packet(payload)
    write_markdown(payload)
    write_docx(payload)
    print(OUT_MD)
    print(OUT_CSV)
    print(OUT_JSON)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()

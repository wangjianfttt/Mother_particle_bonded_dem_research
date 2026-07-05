#!/usr/bin/env python3
"""Check the Computational Particle Mechanics upload package."""

from __future__ import annotations

import csv
import hashlib
import io
import json
import os
import re
import subprocess
import sys
import zipfile
from pathlib import Path

try:
    from docx import Document
except ImportError as exc:  # pragma: no cover - environment message
    raise SystemExit(
        "python-docx is required. Use the bundled Codex Python runtime for this check."
    ) from exc


ROOT = Path(__file__).resolve().parents[1]
UPLOAD_DIR = ROOT / "submission_packages" / "computational_particle_mechanics_upload_ready"
UPLOAD_ZIP = ROOT / "submission_packages" / "computational_particle_mechanics_upload_ready.zip"
UPLOAD_SHA = Path(str(UPLOAD_ZIP) + ".sha256")
REPRO_ZIP = ROOT / "submission_packages" / "repaired_submission_package.zip"
REPRO_SHA = ROOT / "submission_packages" / "repaired_submission_package.zip.sha256"
PUBLIC_REPRO_ZIP = (
    ROOT
    / "submission_packages"
    / "computational_particle_mechanics_public_reproducibility_package.zip"
)
PUBLIC_REPRO_SHA = Path(str(PUBLIC_REPRO_ZIP) + ".sha256")
BLINDED_ZIP = ROOT / "submission_packages" / "computational_particle_mechanics_blinded_review_package.zip"
BLINDED_SHA = ROOT / "submission_packages" / "computational_particle_mechanics_blinded_review_package.zip.sha256"
LEGACY_BLINDED_ZIP = ROOT / "submission_packages" / "computational_particle_mechanics_blinded_review_optional.zip"
LEGACY_BLINDED_SHA = ROOT / "submission_packages" / "computational_particle_mechanics_blinded_review_optional.zip.sha256"
CPM_TEX = ROOT / "manuscript" / "computational_particle_mechanics_submission.tex"
COVER_LETTER = ROOT / "manuscript" / "computational_particle_mechanics_cover_letter.md"
CPM_FIELDS = ROOT / "manuscript" / "computational_particle_mechanics_editorial_fields.md"
CPM_HIGHLIGHTS = ROOT / "manuscript" / "computational_particle_mechanics_highlights.md"
OFFICIAL_GUIDE = ROOT / "docs" / "cpm_official_submission_guide_alignment_20260704.md"
GOAL_AUDIT_JSON = ROOT / "docs" / "cpm_goal_completion_audit_20260704.json"
README = ROOT / "README_CPM_SUBMISSION_20260704.md"
ROOT_README = ROOT / "README.md"
CITATION = ROOT / "CITATION.cff"
START_HERE = ROOT / "START_HERE_CPM_SUBMISSION.md"
LIVE_PACKET_DOCX = ROOT / "manuscript" / "computational_particle_mechanics_live_submission_packet.docx"
LIVE_PACKET_MD = ROOT / "docs" / "cpm_live_submission_packet_20260704.md"
LIVE_PACKET_CSV = ROOT / "docs" / "cpm_live_submission_packet_20260704.csv"
LIVE_PACKET_JSON = ROOT / "docs" / "cpm_live_submission_packet_20260704.json"
ACTION_SHEET_MD = ROOT / "docs" / "cpm_live_submission_action_sheet_20260704.md"
ACTION_SHEET_CSV = ROOT / "docs" / "cpm_live_submission_action_sheet_20260704.csv"
ACTION_SHEET_JSON = ROOT / "docs" / "cpm_live_submission_action_sheet_20260704.json"
PDF_QA_JSON = ROOT / "docs" / "cpm_final_pdf_visual_qa_20260704.json"
PDF_QA_MD = ROOT / "docs" / "cpm_final_pdf_visual_qa_20260704.md"
EMAIL_LOOKUP_MD = ROOT / "docs" / "cpm_author_email_public_lookup_20260704.md"
EMAIL_LOOKUP_CSV = ROOT / "docs" / "cpm_author_email_public_lookup_20260704.csv"
EXTERNAL_AUTHOR_ACTIONS = ROOT / "docs" / "cpm_external_author_metadata_final_actions_20260706.md"
SOURCE_DATA_MATRIX = ROOT / "manuscript" / "repaired_full_manuscript_source_data_matrix.csv"
SUPPORT_DOCX = [
    ROOT / "manuscript" / "computational_particle_mechanics_author_email_collection_packet.docx",
    ROOT / "manuscript" / "computational_particle_mechanics_individual_contact_messages.docx",
    ROOT / "manuscript" / "computational_particle_mechanics_coauthor_email_request_zh_en.docx",
    ROOT / "manuscript" / "computational_particle_mechanics_live_submission_checklist.docx",
    LIVE_PACKET_DOCX,
]
SUPPORT_TEXT = [
    ROOT / "manuscript" / "computational_particle_mechanics_author_email_collection_packet.csv",
    ROOT / "manuscript" / "computational_particle_mechanics_author_email_collection_packet.md",
    ROOT / "manuscript" / "computational_particle_mechanics_author_email_collection_packet.txt",
    ROOT / "manuscript" / "computational_particle_mechanics_individual_contact_messages.csv",
    ROOT / "manuscript" / "computational_particle_mechanics_individual_contact_messages.md",
    ROOT / "manuscript" / "computational_particle_mechanics_individual_contact_messages.txt",
    ROOT / "manuscript" / "computational_particle_mechanics_coauthor_email_request_zh_en.txt",
    ROOT / "manuscript" / "computational_particle_mechanics_live_submission_checklist.md",
    LIVE_PACKET_MD,
    LIVE_PACKET_CSV,
    ACTION_SHEET_MD,
    ACTION_SHEET_CSV,
    ACTION_SHEET_JSON,
    PDF_QA_JSON,
    PDF_QA_MD,
    EMAIL_LOOKUP_MD,
    EMAIL_LOOKUP_CSV,
    EXTERNAL_AUTHOR_ACTIONS,
]

EXPECTED_UPLOAD_FILES = {
    "00_title_page_author_details.docx",
    "01_review_manuscript_blinded.pdf",
    "01_review_manuscript_blinded.tex",
    "02_highlights.docx",
    "03_graphical_abstract.pdf",
    "03_graphical_abstract.png",
    "03_graphical_abstract.svg",
    "03_graphical_abstract.tiff",
    "04_declaration_of_competing_interest.docx",
    "05_cover_letter.docx",
    "06_author_emails_and_contributions.docx",
    "07_latex_source.zip",
    "08_editorial_submission_fields.docx",
    "09_main_figures.zip",
    "10_author_email_completion_sheet.docx",
    "10_author_email_completion_sheet.csv",
    "11_full_author_manuscript_for_production.pdf",
    "12_full_author_manuscript_single_column.docx",
    "MANIFEST.csv",
    "README_upload_roles.txt",
}

DOCX_FILES = [
    "00_title_page_author_details.docx",
    "02_highlights.docx",
    "04_declaration_of_competing_interest.docx",
    "05_cover_letter.docx",
    "06_author_emails_and_contributions.docx",
    "08_editorial_submission_fields.docx",
    "10_author_email_completion_sheet.docx",
    "12_full_author_manuscript_single_column.docx",
]

BLINDED_FORBIDDEN_TERMS = [
    "Acknowledgements",
    "Acknowledgments",
    "Author contributions",
    "Funding",
    "Jian Wang",
    "Siyu Wang",
    "Hang Zhang",
    "Ming-Zhun Lei",
    "Wei Wen",
    "Qi-Gang Wu",
    "Gang Shen",
    "Haishun Deng",
    "wjfttt@mail.ustc.edu.cn",
    "Anhui University of Science and Technology",
    "Institute of Plasma Physics",
    "Chinese Academy of Sciences",
    "Anhui Provincial Natural Science Foundation",
    "DSJJ-2025-08",
    "AIMTEERC202307",
    "China Post-doctoral Science Foundation",
    "2024M753266",
    "2022AH010052",
    "2021yjrc51",
    "2019HSC-CIP006",
    "10.5281/zenodo.20687351",
    "wangjianfttt",
]

FORBIDDEN_READER_TERMS = re.compile(
    r"\b(audit|diagnostic|gate|however|therefore|rather than|not only|CAL|SP-00|PB-00|"
    r"Advanced Powder|Journal of Nuclear Materials|Nuclear Fusion)\b",
    re.IGNORECASE,
)

STALE_200_PARTICLE_WORDING = re.compile(
    r"one 200-particle intact scale check|"
    r"one 200-particle scale-check|"
    r"one 200-mother|"
    r"four zero-loss controls, including a 200-particle intact scale check|"
    r"endpoint table contains four cracking cases, three 100-particle zero-loss controls and one 200-particle|"
    r"eight endpoint mechanism cases|"
    r"completed endpoint dataset contains four localized cracking sequences, four zero-loss controls",
    re.IGNORECASE,
)

AMBIGUOUS_PARENT_PARTICLE_WORDING = re.compile(
    r"\b(?:100|200)-particle\b",
    re.IGNORECASE,
)


def fail(message: str) -> None:
    raise SystemExit(f"FAIL: {message}")


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def check_sha_file(zip_path: Path, sha_path: Path) -> None:
    if not zip_path.exists():
        fail(f"missing zip: {zip_path}")
    if not sha_path.exists():
        fail(f"missing checksum: {sha_path}")
    expected = sha_path.read_text(encoding="utf-8").split()[0]
    actual = sha256(zip_path)
    if actual != expected:
        fail(f"checksum mismatch for {zip_path.name}")


def check_zip(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as zf:
        bad = zf.testzip()
        if bad is not None:
            fail(f"bad zip member in {path.name}: {bad}")
        return zf.namelist()


def check_manifest() -> list[dict[str, str]]:
    manifest = UPLOAD_DIR / "MANIFEST.csv"
    if not manifest.exists():
        fail("missing upload manifest")
    rows = list(csv.DictReader(manifest.open(encoding="utf-8")))
    if len(rows) != 19:
        fail(f"expected 19 upload manifest rows, found {len(rows)}")
    paths = {row["path"] for row in rows}
    required = EXPECTED_UPLOAD_FILES - {"MANIFEST.csv"}
    missing = sorted(required - paths)
    if missing:
        fail(f"manifest missing files: {missing}")
    for row in rows:
        path = UPLOAD_DIR / row["path"]
        if not path.exists():
            fail(f"manifest path missing on disk: {row['path']}")
        if str(path.stat().st_size) != row["bytes"]:
            fail(f"manifest byte count mismatch: {row['path']}")
        if sha256(path) != row["sha256"]:
            fail(f"manifest checksum mismatch: {row['path']}")
    return rows


def check_docx_files() -> None:
    for name in DOCX_FILES:
        Document(UPLOAD_DIR / name)


def check_nested_zips() -> None:
    latex_members = check_zip(UPLOAD_DIR / "07_latex_source.zip")
    if len(latex_members) != 9:
        fail(f"expected 9 LaTeX-source members, found {len(latex_members)}")
    if "manuscript/computational_particle_mechanics_submission.tex" not in latex_members:
        fail("LaTeX source zip missing target-specific TeX")

    figure_members = check_zip(UPLOAD_DIR / "09_main_figures.zip")
    if len(figure_members) != 19:
        fail(f"expected 19 main-figure members, found {len(figure_members)}")
    for stem in [
        "figures/main/fig1_workflow",
        "figures/apt_redesign/fig2_single_pebble_template_validation",
        "figures/apt_redesign/fig3_entry_state_validation",
        "figures/apt_redesign/fig4_pilot_fracture_event_sequence",
        "figures/pb007/pb007_replicate_comparison",
        "figures/pb007/pb007_material_strength_response",
    ]:
        for ext in [".pdf", ".png", ".svg"]:
            if stem + ext not in figure_members:
                fail(f"main figure zip missing {stem + ext}")


def split_semicolon_paths(value: str) -> list[str]:
    return [part.strip() for part in value.split(";") if part.strip()]


def check_main_figure_source_package() -> None:
    """Require each main figure to be source-data-backed and editable."""
    if not SOURCE_DATA_MATRIX.exists():
        fail("missing repaired full manuscript source-data matrix")

    rows = list(csv.DictReader(SOURCE_DATA_MATRIX.open(encoding="utf-8")))
    by_item = {row["display_item"]: row for row in rows}
    expected_items = {f"Fig. {idx}" for idx in range(1, 7)}
    missing_items = sorted(expected_items - set(by_item))
    if missing_items:
        fail(f"source-data matrix missing main figure rows: {missing_items}")

    figure_members = set(check_zip(UPLOAD_DIR / "09_main_figures.zip"))
    for item in sorted(expected_items, key=lambda text: int(text.split()[-1])):
        row = by_item[item]
        output_files = split_semicolon_paths(row.get("output_files", ""))
        source_files = split_semicolon_paths(row.get("source_data", ""))
        scripts = split_semicolon_paths(row.get("regeneration_or_check_script", ""))
        if not output_files:
            fail(f"{item} has no output files in source-data matrix")
        if not source_files:
            fail(f"{item} has no source data or schematic generation source")
        if not scripts:
            fail(f"{item} has no regeneration/check script")

        stems = {str(Path(rel).with_suffix("")) for rel in output_files}
        if len(stems) != 1:
            fail(f"{item} output files do not share one figure stem: {output_files}")
        stem = next(iter(stems))
        for ext in [".pdf", ".png", ".svg"]:
            rel = stem + ext
            if rel not in output_files:
                fail(f"{item} source-data matrix missing {rel}")
            if rel not in figure_members:
                fail(f"{item} main figure zip missing {rel}")
        tiff_rel = stem + ".tiff"
        if not (ROOT / tiff_rel).exists():
            fail(f"{item} missing local high-resolution TIFF: {tiff_rel}")

        for rel in output_files + source_files + scripts:
            if not (ROOT / rel).exists():
                fail(f"{item} references missing source-package file: {rel}")

    readme_members = [member for member in figure_members if member.endswith("README_main_figures.txt")]
    expected_member_count = 1 + 6 * 3
    if len(figure_members) != expected_member_count or len(readme_members) != 1:
        fail(
            "main figure zip should contain one README plus PDF/PNG/SVG for six figures"
        )


def check_author_email_sheet() -> None:
    sheet = UPLOAD_DIR / "10_author_email_completion_sheet.csv"
    rows = list(csv.DictReader(sheet.open(encoding="utf-8")))
    if len(rows) != 8:
        fail(f"expected 8 author rows, found {len(rows)}")
    missing = [row for row in rows if row["Status"] == "Missing"]
    if len(missing) != 7:
        fail(f"expected 7 missing coauthor e-mails, found {len(missing)}")
    if not any(row["Author"] == "Jian Wang" and row["Status"] == "Available" for row in rows):
        fail("corresponding author e-mail is not marked available")


def check_highlights() -> None:
    lines = [
        line.strip("- ").strip()
        for line in CPM_HIGHLIGHTS.read_text(encoding="utf-8").splitlines()
        if line.strip().startswith("-")
    ]
    if len(lines) != 5:
        fail(f"expected 5 highlights, found {len(lines)}")
    too_long = [(idx, len(text), text) for idx, text in enumerate(lines, 1) if len(text) > 85]
    if too_long:
        fail(f"highlight length exceeds 85 characters: {too_long}")


def check_abstract_word_count() -> None:
    text = CPM_TEX.read_text(encoding="utf-8")
    match = re.search(r"\\begin\{abstract\}(.*?)\\end\{abstract\}", text, re.S)
    if not match:
        fail("manuscript TeX missing abstract environment")
    abstract = match.group(1)
    abstract = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?(?:\{([^{}]*)\})?", lambda item: item.group(1) or "", abstract)
    abstract = re.sub(r"[{}$]", " ", abstract)
    words = re.findall(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)?", abstract)
    if len(words) > 250:
        fail(f"abstract exceeds CPM 250-word guide limit: {len(words)} words")


def check_reader_text() -> None:
    for path in [
        ROOT_README,
        CPM_TEX,
        COVER_LETTER,
        CPM_FIELDS,
        CPM_HIGHLIGHTS,
        UPLOAD_DIR / "README_upload_roles.txt",
        README,
    ]:
        text = path.read_text(encoding="utf-8", errors="ignore")
        match = FORBIDDEN_READER_TERMS.search(text)
        if match:
            fail(f"reader-facing residue {match.group(0)!r} in {path}")


def check_repository_metadata() -> None:
    if not ROOT_README.exists():
        fail("missing root README.md for GitHub repository")
    if not CITATION.exists():
        fail("missing CITATION.cff for GitHub/Zenodo citation metadata")

    readme = ROOT_README.read_text(encoding="utf-8", errors="ignore")
    for required in [
        "Bonded-template DEM reveals strength- and topology-dependent fracture-event sequences in packed brittle ceramic pebbles",
        "Computational Particle Mechanics",
        "Elsevier/ScienceDirect double-anonymized",
        "https://doi.org/10.5281/zenodo.20687351",
        "submission_packages/computational_particle_mechanics_upload_ready.zip",
        "submission_packages/repaired_submission_package.zip",
        "ready_for_live_submission_after_external_metadata",
        "CITATION.cff",
    ]:
        if required not in readme:
            fail(f"root README missing current repository metadata: {required}")
    for forbidden in [
        "Journal of Nuclear Materials manuscript",
        "Acceptance-gated bonded-template DEM reveals localized fracture sequences in Li4SiO4 ceramic breeder beds",
        "Current Milestone",
        "第一阶段目标",
        "APT manuscript",
    ]:
        if forbidden in readme:
            fail(f"root README retains stale manuscript-route text: {forbidden}")

    citation = CITATION.read_text(encoding="utf-8", errors="ignore")
    for required in [
        "cff-version: 1.2.0",
        "doi: 10.5281/zenodo.20687351",
        "repository-code: \"https://github.com/wangjianfttt/Mother_particle_bonded_dem_research\"",
        "family-names: Wang",
        "given-names: Jian",
    ]:
        if required not in citation:
            fail(f"CITATION.cff missing metadata: {required}")


def check_stale_200_particle_wording() -> None:
    text_paths = [
        CPM_TEX,
        COVER_LETTER,
        CPM_FIELDS,
        LIVE_PACKET_MD,
        LIVE_PACKET_CSV,
        LIVE_PACKET_JSON,
        README,
        START_HERE,
        ROOT / "docs" / "cpm_reviewer_risk_preflight_20260704.md",
        ROOT / "docs" / "cpm_literature_gap_map_20260704.md",
        ROOT / "manuscript" / "repaired_manuscript_claim_evidence_matrix.csv",
        ROOT / "manuscript" / "repaired_full_manuscript_source_data_matrix.csv",
        ROOT / "manuscript" / "computational_particle_mechanics_blinded_submission.tex",
        UPLOAD_DIR / "01_review_manuscript_blinded.tex",
    ]
    for path in text_paths:
        if not path.exists():
            fail(f"missing stale-wording scan target: {path}")
        text = path.read_text(encoding="utf-8", errors="ignore")
        match = STALE_200_PARTICLE_WORDING.search(text)
        if match:
            fail(f"stale 200-particle wording {match.group(0)!r} in {path}")
        match = AMBIGUOUS_PARENT_PARTICLE_WORDING.search(text)
        if match:
            fail(f"ambiguous parent-particle wording {match.group(0)!r} in {path}")

    docx_paths = SUPPORT_DOCX + [UPLOAD_DIR / name for name in DOCX_FILES]
    for path in docx_paths:
        if not path.exists():
            fail(f"missing stale-wording DOCX scan target: {path}")
        text = docx_text(path)
        match = STALE_200_PARTICLE_WORDING.search(text)
        if match:
            fail(f"stale 200-particle wording {match.group(0)!r} in DOCX {path}")
        match = AMBIGUOUS_PARENT_PARTICLE_WORDING.search(text)
        if match:
            fail(f"ambiguous parent-particle wording {match.group(0)!r} in DOCX {path}")

    with zipfile.ZipFile(UPLOAD_ZIP) as zf:
        for name in zf.namelist():
            payload = zf.read(name)
            if name.endswith(".docx"):
                text = docx_text(io.BytesIO(payload))
            elif name.endswith((".txt", ".md", ".csv", ".tex", ".json")):
                text = payload.decode("utf-8", errors="ignore")
            else:
                continue
            match = STALE_200_PARTICLE_WORDING.search(text)
            if match:
                fail(f"stale 200-particle wording {match.group(0)!r} in {UPLOAD_ZIP.name}:{name}")
            match = AMBIGUOUS_PARENT_PARTICLE_WORDING.search(text)
            if match:
                fail(f"ambiguous parent-particle wording {match.group(0)!r} in {UPLOAD_ZIP.name}:{name}")

    with zipfile.ZipFile(BLINDED_ZIP) as zf:
        for name in zf.namelist():
            payload = zf.read(name)
            if name.endswith(".docx"):
                text = docx_text(io.BytesIO(payload))
            elif name.endswith((".txt", ".md", ".csv", ".tex", ".json")):
                text = payload.decode("utf-8", errors="ignore")
            else:
                continue
            match = STALE_200_PARTICLE_WORDING.search(text)
            if match:
                fail(f"stale 200-particle wording {match.group(0)!r} in {BLINDED_ZIP.name}:{name}")
            match = AMBIGUOUS_PARENT_PARTICLE_WORDING.search(text)
            if match:
                fail(f"ambiguous parent-particle wording {match.group(0)!r} in {BLINDED_ZIP.name}:{name}")


def docx_text(source: Path | io.BytesIO) -> str:
    doc = Document(source)
    parts = [para.text for para in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def check_doi_and_target() -> None:
    text = CPM_TEX.read_text(encoding="utf-8")
    if "10.5281/zenodo.20687351" not in text:
        fail("manuscript TeX missing Zenodo DOI")
    if r"\journal{Computational Particle Mechanics}" not in text:
        fail("manuscript TeX missing target journal")
    fields = CPM_FIELDS.read_text(encoding="utf-8")
    if fields.count("10.5281/zenodo.20687351") < 2:
        fail("editorial fields missing DOI statements")


def check_support_docs() -> None:
    if not START_HERE.exists():
        fail("missing START_HERE_CPM_SUBMISSION.md")
    for path in SUPPORT_DOCX:
        if not path.exists():
            fail(f"missing support DOCX: {path.name}")
        Document(path)
    for path in SUPPORT_TEXT:
        if not path.exists() or path.stat().st_size == 0:
            fail(f"missing or empty support text file: {path.name}")
    external_author_actions = EXTERNAL_AUTHOR_ACTIONS.read_text(encoding="utf-8")
    for term in [
        "Ask directly",
        "Confirm public candidates before use",
        "Siyu Wang",
        "Hang Zhang",
        "Qi-Gang Wu",
        "leimz@ipp.ac.cn",
        "wenwei@ipp.ac.cn",
        "shenganghit@163.com",
        "269469122@qq.com",
        "Preview the generated submission PDF before final submission",
    ]:
        if term not in external_author_actions:
            fail(f"external author metadata final-actions note missing term: {term}")
    start = START_HERE.read_text(encoding="utf-8")
    for required in [
        "computational_particle_mechanics_upload_ready.zip",
        "00_title_page_author_details.docx",
        "01_review_manuscript_blinded.pdf",
        "10_author_email_completion_sheet.docx",
        "12_full_author_manuscript_single_column.docx",
        "computational_particle_mechanics_author_email_collection_packet.docx",
        "computational_particle_mechanics_individual_contact_messages.docx",
        "computational_particle_mechanics_blinded_review_package.zip",
        "Use the blinded manuscript as the review manuscript",
        "computational_particle_mechanics_live_submission_packet.docx",
        "cpm_live_submission_packet_20260704.md",
        "cpm_live_submission_action_sheet_20260704.md",
        "cpm_final_pdf_visual_qa_20260704.md",
        "cpm_author_email_public_lookup_20260704.md",
        "cpm_official_submission_guide_alignment_20260704.md",
        "scripts/check_computational_particle_mechanics_submission_package.py",
        "10.5281/zenodo.20687351",
        "Four public candidate e-mail records",
        "confirmation aids only",
        "Reduced reproducibility package CPM support members: `43/43` present",
        "Current final PDF visual QA: `PASS`, 18 pages, 0 blank pages, author-production PDF SHA match, blinded review PDF checked",
        "Elsevier/ScienceDirect",
        "double-anonymized review",
        "author-bearing PDF/Word backup",
    ]:
        if required not in start:
            fail(f"START_HERE missing {required}")
    for forbidden in [
        "Springer/SNAPP",
        "single-blind",
        "live-system branch",
        "live submit button",
    ]:
        if forbidden in start:
            fail(f"START_HERE retains stale submission-route wording: {forbidden}")
    if not LIVE_PACKET_JSON.exists():
        fail("missing live-submission packet JSON")
    if not ACTION_SHEET_JSON.exists():
        fail("missing live-submission action-sheet JSON")
    readiness_json = ROOT / "docs" / "cpm_submission_readiness_report_20260704.json"
    if not readiness_json.exists():
        fail("missing CPM submission readiness JSON")
    readiness_payload = json.loads(readiness_json.read_text(encoding="utf-8"))
    coauthor_request = (
        ROOT / "manuscript" / "computational_particle_mechanics_coauthor_email_request_zh_en.txt"
    ).read_text(encoding="utf-8")
    for term in [
        "leimz@ipp.ac.cn",
        "wenwei@ipp.ac.cn",
        "shenganghit@163.com",
        "269469122@qq.com",
        "未确认前直接填入投稿系统",
        "before author confirmation",
    ]:
        if term not in coauthor_request:
            fail(f"coauthor e-mail request missing confirmation term: {term}")
    collection_packet = (
        ROOT / "manuscript" / "computational_particle_mechanics_author_email_collection_packet.md"
    ).read_text(encoding="utf-8")
    for term in [
        "Copy-ready short message",
        "Siyu Wang",
        "Hang Zhang",
        "Qi-Gang Wu",
        "Public candidates are confirmation aids only",
    ]:
        if term not in collection_packet:
            fail(f"author e-mail collection packet missing term: {term}")
    individual_messages = (
        ROOT / "manuscript" / "computational_particle_mechanics_individual_contact_messages.md"
    ).read_text(encoding="utf-8")
    for term in [
        "CPM individual coauthor contact messages",
        "confirm_public_candidate",
        "ask_directly",
        "Siyu Wang",
        "Qi-Gang Wu",
        "leimz@ipp.ac.cn",
        "wenwei@ipp.ac.cn",
        "shenganghit@163.com",
        "269469122@qq.com",
    ]:
        if term not in individual_messages:
            fail(f"individual contact messages missing term: {term}")
    payload = json.loads(LIVE_PACKET_JSON.read_text(encoding="utf-8"))
    for key, value in [
        ("target_journal", "Computational Particle Mechanics"),
        ("submission_route", "Elsevier / ScienceDirect double-anonymized route"),
        ("upload_package", "submission_packages/computational_particle_mechanics_upload_ready.zip"),
        ("blinded_review_package", "submission_packages/computational_particle_mechanics_blinded_review_package.zip"),
        ("legacy_blinded_package", "submission_packages/computational_particle_mechanics_blinded_review_optional.zip"),
        ("reduced_reproducibility_package", "submission_packages/repaired_submission_package.zip"),
        ("missing_email_count", 7),
    ]:
        if payload.get(key) != value:
            fail(f"live-submission packet JSON mismatch for {key}")
    if payload.get("candidate_email_lookup") != "docs/cpm_author_email_public_lookup_20260704.md":
        fail("live-submission packet JSON missing candidate e-mail lookup")
    if payload.get("candidate_email_count") != 4:
        fail("expected four public candidate e-mails for confirmation")
    action_payload = json.loads(ACTION_SHEET_JSON.read_text(encoding="utf-8"))
    if action_payload.get("missing_email_count") != 7:
        fail("live-submission action sheet does not preserve missing e-mail count")
    if action_payload.get("upload_package_sha256") != readiness_payload.get("upload_package_sha256"):
        fail("live-submission action sheet upload-package SHA is stale")
    if "reduced_reproducibility_package_sha256" in action_payload:
        fail("live-submission action sheet contains a self-referential reduced-package SHA")
    if "readiness_report" not in action_payload.get(
        "reduced_reproducibility_package_sha256_source", ""
    ):
        fail("live-submission action sheet missing reduced-package SHA source note")
    actions = action_payload.get("actions", [])
    if len(actions) != 16:
        fail(f"expected 16 live-submission action rows, found {len(actions)}")
    if any(row.get("file_status") != "present" for row in actions if isinstance(row, dict)):
        fail("live-submission action sheet includes a missing file")
    action_text = ACTION_SHEET_MD.read_text(encoding="utf-8")
    for term in [
        "CPM Live Submission Action Sheet",
        "Preview the system-generated PDF before final submit",
        "Review manuscript for Elsevier/ScienceDirect double-anonymized route",
        "required_for_double_anonymized_review",
        "production_source_or_explicit_system_request",
        "formal_upload_package",
    ]:
        if term == "formal_upload_package":
            if term not in ACTION_SHEET_CSV.read_text(encoding="utf-8"):
                fail(f"live-submission action sheet CSV missing {term}")
        elif term not in action_text:
            fail(f"live-submission action sheet missing {term}")
    final_actions = payload.get("final_external_actions", [])
    if not isinstance(final_actions, list):
        fail("live-submission packet final actions are not a list")
    if len(final_actions) != len(set(final_actions)):
        fail("live-submission packet final actions contain exact duplicates")
    article_type_actions = [
        item for item in final_actions if isinstance(item, str) and "article type/category" in item
    ]
    if len(article_type_actions) != 1:
        fail("live-submission packet should contain exactly one article type/category action")
    if "reduced reproducibility package SHA256" in action_text:
        fail("live-submission action sheet should defer reduced package SHA to readiness report")
    if not PDF_QA_JSON.exists():
        fail("missing CPM final PDF visual QA JSON")
    pdf_qa = json.loads(PDF_QA_JSON.read_text(encoding="utf-8"))
    expected_pdf_qa = {
        "status": "PASS",
        "page_count": 18,
        "author_production_matches_manuscript_pdf": True,
        "blank_page_count": 0,
        "contains_title": True,
        "contains_doi": True,
        "contains_references_heading": True,
        "unresolved_reference_tokens": 0,
    }
    for key, expected in expected_pdf_qa.items():
        if pdf_qa.get(key) != expected:
            fail(f"CPM final PDF visual QA mismatch for {key}: {pdf_qa.get(key)!r}")
    current_pdf_shas = {
        "manuscript_pdf_sha256": sha256(CPM_TEX.with_suffix(".pdf")),
        "author_production_pdf_sha256": sha256(
            UPLOAD_DIR / "11_full_author_manuscript_for_production.pdf"
        ),
        "blinded_review_pdf_sha256": sha256(UPLOAD_DIR / "01_review_manuscript_blinded.pdf"),
    }
    for key, current_sha in current_pdf_shas.items():
        if pdf_qa.get(key) != current_sha:
            fail(
                "CPM final PDF visual QA is stale for "
                f"{key}: report={pdf_qa.get(key)!r}, current={current_sha!r}"
            )
    contact_sheet = ROOT / pdf_qa.get("contact_sheet", "")
    if not contact_sheet.exists() or contact_sheet.stat().st_size == 0:
        fail("CPM final PDF contact sheet is missing or empty")
    if pdf_qa.get("blinded_review_forbidden_hits") != []:
        fail("CPM final PDF QA reports identifying terms in blinded review manuscript")


def check_official_guide_alignment() -> None:
    if not OFFICIAL_GUIDE.exists():
        fail("missing official submission-guide alignment report")
    text = OFFICIAL_GUIDE.read_text(encoding="utf-8")
    required = [
        "ScienceDirect Guide for Authors",
        "Springer journal transition notice",
        "Elsevier/ScienceDirect",
        "Review route",
        "Double-anonymized review",
        "Abstract length",
        "not exceed 250 words",
        "required_for_double_anonymized_review",
        "review manuscript",
        "computational_particle_mechanics_blinded_review_package.zip",
        "external_metadata_pending",
    ]
    for term in required:
        if term not in text:
            fail(f"official guide alignment report missing {term!r}")
    forbidden = [
        "highlights are required for this journal family workflow",
        "Springer/SNAPP",
        "single-blind",
        "live-system branch",
    ]
    for term in forbidden:
        if term in text:
            fail(f"official guide alignment retains unsupported requirement wording: {term}")


def check_goal_completion_audit() -> None:
    if not GOAL_AUDIT_JSON.exists():
        fail("missing goal-level completion audit JSON")
    payload = json.loads(GOAL_AUDIT_JSON.read_text(encoding="utf-8"))
    if payload.get("overall_status") != "ready_after_external_author_metadata":
        fail("goal-level completion audit does not record the correct external-metadata boundary")
    rows = payload.get("rows", [])
    statuses = {row.get("requirement"): row.get("status") for row in rows if isinstance(row, dict)}
    if statuses.get("External author metadata for live system") != "external_pending":
        fail("goal-level completion audit missing external author-metadata pending status")
    if payload.get("large_raw_residue_count") != 0:
        fail("goal-level completion audit reports remaining large raw dump/restart residues")


def check_blinded_review_package() -> None:
    check_sha_file(BLINDED_ZIP, BLINDED_SHA)
    names = check_zip(BLINDED_ZIP)
    required = {
        "computational_particle_mechanics_blinded_review_package/01_blinded_manuscript.pdf",
        "computational_particle_mechanics_blinded_review_package/01_blinded_manuscript.tex",
        "computational_particle_mechanics_blinded_review_package/README_blinded_review.txt",
        "computational_particle_mechanics_blinded_review_package/references.bib",
    }
    missing = sorted(required - set(names))
    if missing:
        fail(f"blinded-review zip missing files: {missing}")
    with zipfile.ZipFile(BLINDED_ZIP) as zf:
        tex = zf.read("computational_particle_mechanics_blinded_review_package/01_blinded_manuscript.tex").decode("utf-8", errors="ignore")
    found = [term for term in BLINDED_FORBIDDEN_TERMS if term in tex]
    if found:
        fail(f"blinded manuscript source still contains identifying terms: {found}")
    upload_blinded_tex = UPLOAD_DIR / "01_review_manuscript_blinded.tex"
    upload_found = [
        term
        for term in BLINDED_FORBIDDEN_TERMS
        if term in upload_blinded_tex.read_text(encoding="utf-8", errors="ignore")
    ]
    if upload_found:
        fail(f"upload blinded manuscript source still contains identifying terms: {upload_found}")
    check_sha_file(LEGACY_BLINDED_ZIP, LEGACY_BLINDED_SHA)
    check_zip(LEGACY_BLINDED_ZIP)


def check_scientific_alignment() -> None:
    subprocess.run(
        [sys.executable, str(ROOT / "scripts" / "check_cpm_scientific_alignment.py")],
        cwd=ROOT,
        check=True,
    )


def check_reviewer_risk_preflight() -> None:
    subprocess.run(
        [sys.executable, str(ROOT / "scripts" / "check_cpm_reviewer_risk_preflight.py")],
        cwd=ROOT,
        check=True,
    )


def check_public_reproducibility_package() -> None:
    check_sha_file(PUBLIC_REPRO_ZIP, PUBLIC_REPRO_SHA)
    subprocess.run(
        [sys.executable, str(ROOT / "scripts" / "check_cpm_public_repro_package.py")],
        cwd=ROOT,
        check=True,
    )


def main() -> None:
    check_sha_file(UPLOAD_ZIP, UPLOAD_SHA)
    check_sha_file(REPRO_ZIP, REPRO_SHA)
    check_public_reproducibility_package()
    check_zip(UPLOAD_ZIP)
    check_manifest()
    check_docx_files()
    check_nested_zips()
    check_main_figure_source_package()
    check_author_email_sheet()
    check_highlights()
    check_abstract_word_count()
    check_reader_text()
    check_repository_metadata()
    check_stale_200_particle_wording()
    check_doi_and_target()
    check_support_docs()
    check_official_guide_alignment()
    if os.environ.get("CPM_SKIP_GOAL_AUDIT") != "1":
        check_goal_completion_audit()
    check_blinded_review_package()
    check_scientific_alignment()
    check_reviewer_risk_preflight()
    print(
        "PASS CPM submission package: manifest=19, figures=19, "
        "figure source-data gate, docx=11, DOI, guide alignment, public "
        "reproducibility package, live packet and double-anonymous blinded "
        "review manuscript verified"
    )


if __name__ == "__main__":
    main()

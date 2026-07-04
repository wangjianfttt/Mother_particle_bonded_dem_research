#!/usr/bin/env python3
"""Build practical support documents for CPM online submission."""

from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "manuscript"

TITLE = (
    "Bonded-template DEM reveals strength- and topology-dependent "
    "fracture-event sequences in packed brittle ceramic pebbles"
)

MISSING_AUTHORS = [
    "Siyu Wang",
    "Hang Zhang",
    "Ming-Zhun Lei",
    "Wei Wen",
    "Qi-Gang Wu",
    "Gang Shen",
    "Haishun Deng",
]

PUBLIC_CANDIDATE_EMAILS = [
    (
        "Ming-Zhun Lei",
        "leimz@ipp.ac.cn",
        "请雷明准老师确认该邮箱是否可用于投稿系统",
        "Please confirm whether this e-mail may be used in the submission system.",
    ),
    (
        "Wei Wen",
        "wenwei@ipp.ac.cn",
        "请文伟老师确认该邮箱是否可用于投稿系统",
        "Please confirm whether this e-mail may be used in the submission system.",
    ),
    (
        "Gang Shen",
        "shenganghit@163.com",
        "请沈刚老师确认该邮箱是否可用于投稿系统",
        "Please confirm whether this e-mail may be used in the submission system.",
    ),
    (
        "Haishun Deng",
        "269469122@qq.com",
        "请邓海顺老师确认该邮箱是否可用于投稿系统",
        "Please confirm whether this e-mail may be used in the submission system.",
    ),
]

AUTHOR_EMAIL_ACTION_ROWS = [
    {
        "author": "Siyu Wang",
        "affiliation": "Anhui University of Science and Technology",
        "current_status": "Missing",
        "public_candidate": "",
        "source": "No reliable public candidate after local search and targeted web search.",
        "action": "Ask author directly for preferred submission-system e-mail.",
    },
    {
        "author": "Hang Zhang",
        "affiliation": "Anhui University of Science and Technology",
        "current_status": "Missing",
        "public_candidate": "",
        "source": "No reliable public candidate after local search and targeted web search.",
        "action": "Ask author directly for preferred submission-system e-mail.",
    },
    {
        "author": "Ming-Zhun Lei",
        "affiliation": "Institute of Plasma Physics, Chinese Academy of Sciences",
        "current_status": "Missing",
        "public_candidate": "leimz@ipp.ac.cn",
        "source": "https://pmc.ncbi.nlm.nih.gov/articles/PMC8234535/",
        "action": "Ask author to confirm whether this public candidate may be used.",
    },
    {
        "author": "Wei Wen",
        "affiliation": "Anhui University of Science and Technology; Institute of Plasma Physics, Chinese Academy of Sciences",
        "current_status": "Missing",
        "public_candidate": "wenwei@ipp.ac.cn",
        "source": "https://www.ipp.cas.cn/yjs/xcsc/202207/P020250115777706650358.pdf",
        "action": "Ask author to confirm whether this public candidate may be used.",
    },
    {
        "author": "Qi-Gang Wu",
        "affiliation": "Institute of Plasma Physics, Chinese Academy of Sciences",
        "current_status": "Missing",
        "public_candidate": "",
        "source": "No reliable personal e-mail candidate after local search and targeted web search.",
        "action": "Ask author directly for preferred submission-system e-mail.",
    },
    {
        "author": "Gang Shen",
        "affiliation": "Anhui University of Science and Technology",
        "current_status": "Missing",
        "public_candidate": "shenganghit@163.com",
        "source": "https://journals.sagepub.com/doi/10.1177/09544062241299521",
        "action": "Ask author to confirm whether this public candidate may be used.",
    },
    {
        "author": "Haishun Deng",
        "affiliation": "Anhui University of Science and Technology",
        "current_status": "Missing",
        "public_candidate": "269469122@qq.com",
        "source": "https://pmc.ncbi.nlm.nih.gov/articles/PMC11637116/",
        "action": "Ask author to confirm whether this public candidate may be used.",
    },
]


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)


def add_title(doc: Document, heading: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(heading)
    r.bold = True
    r.font.size = Pt(15)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(TITLE)
    r2.italic = True
    r2.font.size = Pt(9)


def write_email_request() -> None:
    subject = "请补充CPM投稿作者邮箱 / Author e-mail confirmation for CPM submission"
    body_zh = [
        "各位老师、同学好，",
        "",
        "Li4SiO4 bonded-template DEM 颗粒破碎论文准备转投 Computational Particle Mechanics。投稿系统可能要求填写每位作者的邮箱地址。当前稿件中通讯作者邮箱已经确认：Jian Wang, wjfttt@mail.ustc.edu.cn。",
        "",
        "请下面几位作者各自提供一个投稿系统使用的机构邮箱，或确认可以使用的常用学术邮箱：",
        *[f"- {name}" for name in MISSING_AUTHORS],
        "",
        "我另外查到四条公开候选邮箱，仅供本人确认，不会在未确认前直接填入投稿系统：",
        *[f"- {name}: {email}（{note_zh}）" for name, email, note_zh, _note_en in PUBLIC_CANDIDATE_EMAILS],
        "",
        "同时请确认作者顺序、单位和贡献描述没有问题。如需查看当前记录，可参考随附的 author e-mail completion sheet。",
        "",
        "谢谢！",
        "Jian Wang",
    ]
    body_en = [
        "Dear coauthors,",
        "",
        "We are preparing the Li4SiO4 bonded-template DEM manuscript for submission to Computational Particle Mechanics. The online submission system may require an e-mail address for every author. The corresponding author e-mail has been confirmed as Jian Wang, wjfttt@mail.ustc.edu.cn.",
        "",
        "Please provide one institutional or preferred academic e-mail address for the following authors:",
        *[f"- {name}" for name in MISSING_AUTHORS],
        "",
        "Four public candidate e-mail records have been found for confirmation only. They will not be entered into the submission system before author confirmation:",
        *[f"- {name}: {email} ({note_en})" for name, email, _note_zh, note_en in PUBLIC_CANDIDATE_EMAILS],
        "",
        "Please also confirm that the author order, affiliations and contribution descriptions are correct. The current author e-mail completion sheet is attached for reference.",
        "",
        "Best regards,",
        "Jian Wang",
    ]
    txt = MANUSCRIPT / "computational_particle_mechanics_coauthor_email_request_zh_en.txt"
    txt.write_text(
        f"Subject: {subject}\n\n" + "\n".join(body_zh) + "\n\n---\n\n" + "\n".join(body_en) + "\n",
        encoding="utf-8",
    )

    doc = Document()
    configure(doc)
    add_title(doc, "Coauthor e-mail request template")
    doc.add_heading("Suggested e-mail subject", level=1)
    doc.add_paragraph(subject)
    doc.add_heading("Chinese version", level=1)
    for line in body_zh:
        doc.add_paragraph(line)
    doc.add_heading("English version", level=1)
    for line in body_en:
        doc.add_paragraph(line)
    doc.save(MANUSCRIPT / "computational_particle_mechanics_coauthor_email_request_zh_en.docx")


def write_live_checklist() -> None:
    checklist = [
        ("Preflight", "Run scripts/check_computational_particle_mechanics_submission_package.py and confirm PASS."),
        ("Author e-mails", "Fill seven missing coauthor e-mails if the live system requires all author e-mails."),
        ("Candidate e-mails", "Ask Ming-Zhun Lei, Wei Wen, Gang Shen and Haishun Deng to confirm whether the four public candidate e-mails may be used."),
        ("Article type", "Select Research article or the closest equivalent in the live Elsevier/ScienceDirect system."),
        ("Manuscript", "Upload 01_manuscript.pdf."),
        ("Highlights", "Upload 02_highlights.docx or paste the five Highlights from 08_editorial_submission_fields.docx."),
        ("Graphical abstract", "Upload 03_graphical_abstract.png or 03_graphical_abstract.tiff."),
        ("Declaration", "Upload 04_declaration_of_competing_interest.docx."),
        ("Cover letter", "Upload 05_cover_letter.docx."),
        ("Authors and CRediT", "Use 06_author_emails_and_contributions.docx and the completed e-mail sheet."),
        ("Blinded manuscript if requested", "Use computational_particle_mechanics_blinded_review_optional.zip only if the live system requests a blinded manuscript for double-anonymized review."),
        ("LaTeX source", "Upload 07_latex_source.zip if the initial submission system requests source files."),
        ("Figures", "Upload 09_main_figures.zip or individual figure files if the system separates artwork upload."),
        ("Data availability", "Paste the DOI-backed data availability statement from 08_editorial_submission_fields.docx."),
        ("Code availability", "Paste the GitHub/Zenodo code availability statement from 08_editorial_submission_fields.docx."),
        ("Final check", "Download or preview the generated submission PDF before clicking final submit."),
    ]
    doc = Document()
    configure(doc)
    add_title(doc, "CPM live submission checklist")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Status", "Item", "Action"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for item, action in checklist:
        cells = table.add_row().cells
        cells[0].text = "[ ]"
        cells[1].text = item
        cells[2].text = action
    doc.add_paragraph(
        "Current external item: seven coauthor e-mail addresses are not present in the local records."
    )
    doc.save(MANUSCRIPT / "computational_particle_mechanics_live_submission_checklist.docx")

    md = MANUSCRIPT / "computational_particle_mechanics_live_submission_checklist.md"
    lines = ["# CPM live submission checklist", "", f"Manuscript: {TITLE}", ""]
    for item, action in checklist:
        lines.append(f"- [ ] **{item}:** {action}")
    lines.append("")
    lines.append("Current external item: seven coauthor e-mail addresses are not present in the local records.")
    md.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_author_email_collection_packet() -> None:
    csv_path = MANUSCRIPT / "computational_particle_mechanics_author_email_collection_packet.csv"
    with csv_path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(AUTHOR_EMAIL_ACTION_ROWS[0]), lineterminator="\n")
        writer.writeheader()
        writer.writerows(AUTHOR_EMAIL_ACTION_ROWS)

    short_zh = [
        "各位老师、同学好，CPM投稿系统可能需要每位作者邮箱。请直接回复：",
        "1. 投稿系统使用邮箱；",
        "2. 作者顺序和单位是否确认；",
        "3. 贡献描述是否确认。",
        "已有公开候选邮箱的作者请确认能否使用该邮箱；没有候选邮箱的作者请提供一个常用学术邮箱。",
    ]
    short_en = [
        "Dear coauthors, the CPM submission system may require one e-mail address for every author. Please reply with:",
        "1. preferred submission-system e-mail;",
        "2. confirmation of author order and affiliation;",
        "3. confirmation of the contribution statement.",
        "Authors with public candidate e-mails should confirm whether the candidate can be used. Authors without a candidate should provide a preferred academic e-mail.",
    ]
    md_lines = [
        "# CPM author e-mail collection packet",
        "",
        f"Manuscript: {TITLE}",
        "",
        "## Copy-ready short message, Chinese",
        "",
        *short_zh,
        "",
        "## Copy-ready short message, English",
        "",
        *short_en,
        "",
        "## Action table",
        "",
        "| Author | Affiliation | Status | Public candidate | Action |",
        "| --- | --- | --- | --- | --- |",
    ]
    for row in AUTHOR_EMAIL_ACTION_ROWS:
        md_lines.append(
            f"| {row['author']} | {row['affiliation']} | {row['current_status']} | "
            f"{row['public_candidate'] or 'None'} | {row['action']} |"
        )
    md_lines.extend(
        [
            "",
            "## Rule",
            "",
            "Public candidates are confirmation aids only. Do not enter them into the live submission system until the corresponding author or coauthor confirms the address.",
            "",
        ]
    )
    md_path = MANUSCRIPT / "computational_particle_mechanics_author_email_collection_packet.md"
    txt_path = MANUSCRIPT / "computational_particle_mechanics_author_email_collection_packet.txt"
    md_path.write_text("\n".join(md_lines), encoding="utf-8")
    txt_path.write_text("\n".join(md_lines), encoding="utf-8")

    doc = Document()
    configure(doc)
    add_title(doc, "Author e-mail collection packet")
    doc.add_heading("Copy-ready short message, Chinese", level=1)
    for line in short_zh:
        doc.add_paragraph(line)
    doc.add_heading("Copy-ready short message, English", level=1)
    for line in short_en:
        doc.add_paragraph(line)
    doc.add_heading("Action table", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Author", "Status", "Public candidate", "Action"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for row in AUTHOR_EMAIL_ACTION_ROWS:
        cells = table.add_row().cells
        cells[0].text = row["author"]
        cells[1].text = row["current_status"]
        cells[2].text = row["public_candidate"] or "None"
        cells[3].text = row["action"]
    doc.add_heading("Rule", level=1)
    doc.add_paragraph(
        "Public candidates are confirmation aids only. Do not enter them into the live submission system until the corresponding author or coauthor confirms the address."
    )
    doc.save(MANUSCRIPT / "computational_particle_mechanics_author_email_collection_packet.docx")


def write_individual_contact_messages() -> None:
    rows = []
    for row in AUTHOR_EMAIL_ACTION_ROWS:
        author = row["author"]
        candidate = row["public_candidate"]
        if candidate:
            zh = (
                f"{author}老师您好，CPM投稿系统可能需要每位作者邮箱。"
                f"我查到一个公开候选邮箱：{candidate}。请您确认该邮箱是否可以用于本次投稿系统；"
                "同时请确认作者顺序、单位和贡献描述是否无误。谢谢！"
            )
            en = (
                f"Dear {author}, the CPM submission system may require one e-mail address for every author. "
                f"I found the following public candidate e-mail: {candidate}. "
                "Could you please confirm whether it may be used in the submission system, and whether the author order, affiliation and contribution statement are correct?"
            )
            action_type = "confirm_public_candidate"
        else:
            zh = (
                f"{author}您好，CPM投稿系统可能需要每位作者邮箱。"
                "请您回复一个可用于投稿系统的常用学术邮箱；同时请确认作者顺序、单位和贡献描述是否无误。谢谢！"
            )
            en = (
                f"Dear {author}, the CPM submission system may require one e-mail address for every author. "
                "Could you please provide a preferred academic e-mail address for the submission system, and confirm whether the author order, affiliation and contribution statement are correct?"
            )
            action_type = "ask_directly"
        rows.append(
            {
                "author": author,
                "action_type": action_type,
                "public_candidate": candidate,
                "chinese_message": zh,
                "english_message": en,
            }
        )

    csv_path = MANUSCRIPT / "computational_particle_mechanics_individual_contact_messages.csv"
    with csv_path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0]), lineterminator="\n")
        writer.writeheader()
        writer.writerows(rows)

    md_lines = [
        "# CPM individual coauthor contact messages",
        "",
        f"Manuscript: {TITLE}",
        "",
        "Rule: public candidate e-mails are confirmation aids only. Do not enter them into the live submission system until the corresponding author or coauthor confirms the address.",
        "",
    ]
    for row in rows:
        md_lines.extend(
            [
                f"## {row['author']}",
                "",
                f"- Action type: `{row['action_type']}`",
                f"- Public candidate: `{row['public_candidate'] or 'None'}`",
                "",
                "Chinese message:",
                "",
                row["chinese_message"],
                "",
                "English message:",
                "",
                row["english_message"],
                "",
            ]
        )
    md_path = MANUSCRIPT / "computational_particle_mechanics_individual_contact_messages.md"
    txt_path = MANUSCRIPT / "computational_particle_mechanics_individual_contact_messages.txt"
    md_path.write_text("\n".join(md_lines), encoding="utf-8")
    txt_path.write_text("\n".join(md_lines), encoding="utf-8")

    doc = Document()
    configure(doc)
    add_title(doc, "Individual coauthor contact messages")
    doc.add_paragraph(
        "Public candidate e-mails are confirmation aids only. Do not enter them into the live submission system until the corresponding author or coauthor confirms the address."
    )
    doc.add_paragraph(
        "Use the Chinese message directly for quick confirmation. The English version is included below each author as a backup."
    )
    for row in rows:
        doc.add_heading(row["author"], level=1)
        meta = doc.add_table(rows=0, cols=2)
        meta.style = "Table Grid"
        for label, value in [
            ("Action", row["action_type"]),
            ("Public candidate", row["public_candidate"] or "None"),
        ]:
            cells = meta.add_row().cells
            cells[0].text = label
            cells[1].text = value
        doc.add_paragraph("Chinese message").runs[0].bold = True
        doc.add_paragraph(row["chinese_message"])
        doc.add_paragraph("English backup message").runs[0].bold = True
        doc.add_paragraph(row["english_message"])
    doc.save(MANUSCRIPT / "computational_particle_mechanics_individual_contact_messages.docx")


def main() -> None:
    MANUSCRIPT.mkdir(exist_ok=True)
    write_email_request()
    write_live_checklist()
    write_author_email_collection_packet()
    write_individual_contact_messages()
    print(MANUSCRIPT / "computational_particle_mechanics_coauthor_email_request_zh_en.docx")
    print(MANUSCRIPT / "computational_particle_mechanics_live_submission_checklist.docx")
    print(MANUSCRIPT / "computational_particle_mechanics_author_email_collection_packet.docx")
    print(MANUSCRIPT / "computational_particle_mechanics_individual_contact_messages.docx")


if __name__ == "__main__":
    main()
